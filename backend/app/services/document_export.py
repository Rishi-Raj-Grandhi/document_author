# app/services/document_export.py
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from typing import Any, Dict
import re


def _add_rich_text(paragraph, text: str):
    """Handles **bold**, *italic*, `code` inside text"""
    if not text:
        return

    # Split by markdown patterns (order matters!)
    patterns = [
        (r"`(.*?)`", "code"),
        (r"\*\*(.*?)\*\*", "bold"),
        (r"__(.*?)__", "bold"),
        (r"\*(.*?)\*", "italic"),
        (r"_([^_]+)_", "italic"),
    ]

    parts = [text]
    for pattern, style in patterns:
        new_parts = []
        for part in parts:
            if isinstance(part, str):
                matches = list(re.finditer(pattern, part))
                if not matches:
                    new_parts.append(part)
                    continue
                pos = 0
                for m in matches:
                    if m.start() > pos:
                        new_parts.append(part[pos:m.start()])
                    new_parts.append((m.group(1), style))
                    pos = m.end()
                if pos < len(part):
                    new_parts.append(part[pos:])
            else:
                new_parts.append(part)
        parts = new_parts

    for part in parts:
        if isinstance(part, tuple):
            content, style = part
            run = paragraph.add_run(content)
            if style == "bold":
                run.bold = True
            elif style == "italic":
                run.italic = True
            elif style == "code":
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                shading_elt = OxmlElement('w:shd')
                shading_elt.set(qn('w:fill'), 'f0f0f0')
                run._r.get_or_add_rPr().append(shading_elt)
        else:
            paragraph.add_run(part)


def export_to_word(document_data: Dict[str, Any]) -> BytesIO:
    doc = Document()

    # Default styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Margins
    for section in doc.sections:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)

    doc.core_properties.title = document_data.get("title", "Untitled")

    # Support multiple root keys that LLMs use
    blocks = (
        document_data.get("blocks") or
        document_data.get("sections") or
        document_data.get("content", []) or
        []
    )

    for block in blocks:
        btype = str(block.get("type", "")).lower()

        # Headings
        if btype == "heading":
            level = max(1, min(block.get("level", 2), 9))
            heading_level = 0 if level == 1 else level
            p = doc.add_heading(block.get("text", ""), level=heading_level)
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # Paragraph (with markdown support)
        if btype in ("paragraph", "text", ""):
            text = block.get("text") or block.get("content", "")
            if text:
                p = doc.add_paragraph()
                _add_rich_text(p, str(text))
            continue

        # Bullet lists
        if btype in ("bullet_list", "ul", "bulletlist"):
            for item in block.get("items", []):
                text = item.get("text") if isinstance(item, dict) else str(item)
                p = doc.add_paragraph(style="List Bullet")
                _add_rich_text(p, text)
            continue

        # Numbered lists
        if btype in ("numbered_list", "ol", "numberedlist"):
            for item in block.get("items", []):
                text = item.get("text") if isinstance(item, dict) else str(item)
                p = doc.add_paragraph(style="List Number")
                _add_rich_text(p, text)
            continue

        # Tables
        if btype == "table":
            rows = block.get("rows", [])
            if not rows:
                continue
            cols = len(rows[0].get("cells", [])) if rows else 0
            table = doc.add_table(rows=len(rows), cols=cols, style="Table Grid")
            for r_idx, row in enumerate(rows):
                cells = row.get("cells", [])
                for c_idx, cell in enumerate(cells):
                    cell_text = cell.get("text") if isinstance(cell, dict) else str(cell)
                    cell_obj = table.cell(r_idx, c_idx)
                    cell_obj.text = ""
                    _add_rich_text(cell_obj.paragraphs[0], cell_text)
            continue

        # Code blocks
        if btype == "code":
            p = doc.add_paragraph(block.get("text", ""), style="No Spacing")
            run = p.runs[0]
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'f0f0f0')
            run._r.get_or_add_rPr().append(shading)
            continue

        # Fallback → plain paragraph
        fallback_text = block.get("text") or block.get("content") or str(block)
        if fallback_text.strip():
            p = doc.add_paragraph()
            _add_rich_text(p, fallback_text)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer